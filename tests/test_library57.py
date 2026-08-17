from io import BytesIO

from docx import Document

from backend.library_service import _parse_docx, _split_by_heading_lines


def test_plain_text_chapters_split_into_sections():
    sections = _split_by_heading_lines(
        "CHAPTER 1\n\nFirst chapter body.\n\nCHAPTER 2\n\nSecond chapter body."
    )
    assert [title for title, _ in sections] == ["CHAPTER 1", "CHAPTER 2"]
    assert "First chapter body" in sections[0][1]
    assert "Second chapter body" in sections[1][1]


def test_docx_heading_import_preserves_chapter_structure():
    document = Document()
    document.core_properties.title = "Imported Test Book"
    document.add_heading("Chapter 1", level=1)
    document.add_paragraph("Opening text.")
    document.add_heading("Chapter 2", level=1)
    document.add_paragraph("Second text.")
    stream = BytesIO()
    document.save(stream)

    suggested, sections = _parse_docx(stream.getvalue())

    assert suggested == "Imported Test Book"
    assert len(sections) == 2
    assert sections[0][0] == "Chapter 1"
    assert sections[1][0] == "Chapter 2"
