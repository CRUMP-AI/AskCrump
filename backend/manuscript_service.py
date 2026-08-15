"""Durable long-form manuscript workspaces and KDP-aware exports for Ask Crump 5.4."""
from __future__ import annotations

from dataclasses import dataclass
from datetime import datetime, timezone
from io import BytesIO
import html
import json
import math
import re
from typing import Any
from uuid import uuid4
import zipfile

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
from reportlab.lib.units import inch
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer

from .ai_service import AIService
from .db import SupabaseDB, eq
from .project_service import ProjectService
from .security import normalize_chat_id


KDP_TRIM_SIZES: dict[str, tuple[float, float]] = {
    "5x8": (5.0, 8.0),
    "5.25x8": (5.25, 8.0),
    "5.5x8.5": (5.5, 8.5),
    "6x9": (6.0, 9.0),
    "6.14x9.21": (6.14, 9.21),
    "7x10": (7.0, 10.0),
    "8x10": (8.0, 10.0),
    "8.5x11": (8.5, 11.0),
}

# Upper page-count envelope for black ink on white paper. Other KDP print
# combinations can have lower limits, so this is a preflight warning rather
# than a substitute for the KDP trim/paper calculator.
KDP_BW_WHITE_MAX_PAGES: dict[str, int] = {
    "5x8": 828,
    "5.25x8": 828,
    "5.5x8.5": 828,
    "6x9": 828,
    "6.14x9.21": 828,
    "7x10": 828,
    "8x10": 828,
    "8.5x11": 590,
}

DEFAULT_TARGET_WORDS = 80_000
DEFAULT_CHAPTER_COUNT = 28
MIN_TARGET_WORDS = 20_000
MAX_TARGET_WORDS = 150_000


def _now() -> str:
    return datetime.now(timezone.utc).isoformat()


def _clean(value: Any, limit: int) -> str:
    return " ".join(str(value or "").split()).strip()[:limit]


def word_count(text: str) -> int:
    return len(re.findall(r"\b[\w’'-]+\b", str(text or "")))


def target_words_from_prompt(prompt: str, default: int = DEFAULT_TARGET_WORDS) -> int:
    text = str(prompt or "")
    match = re.search(r"\b(\d{1,3}(?:,\d{3})+)\s*words?\b", text, re.I)
    if match:
        value = int(match.group(1).replace(",", ""))
        return max(MIN_TARGET_WORDS, min(MAX_TARGET_WORDS, value))
    match = re.search(r"\b(\d{2,3})\s*k\s*(?:words?)?\b", text, re.I)
    if match:
        value = int(match.group(1)) * 1000
        return max(MIN_TARGET_WORDS, min(MAX_TARGET_WORDS, value))
    return max(MIN_TARGET_WORDS, min(MAX_TARGET_WORDS, int(default or DEFAULT_TARGET_WORDS)))


def chapter_count_from_prompt(prompt: str, default: int = DEFAULT_CHAPTER_COUNT) -> int:
    match = re.search(r"\b(\d{1,2})\s+(?:chapters?|sections?)\b", str(prompt or ""), re.I)
    value = int(match.group(1)) if match else int(default or DEFAULT_CHAPTER_COUNT)
    return max(8, min(80, value))


def title_from_prompt(prompt: str, fallback: str = "Untitled Manuscript") -> str:
    text = " ".join(str(prompt or "").split())
    match = re.search(
        r"\b(?:called|titled)\s+[\"“']?([^\"”'.,!?]{2,120})",
        text,
        re.I,
    )
    if match:
        return _clean(match.group(1), 180)
    return fallback


def estimate_pages(total_words: int, front_matter_pages: int = 4) -> int:
    body_pages = max(1, math.ceil(max(0, total_words) / 275))
    return max(24, math.ceil((body_pages + front_matter_pages) * 1.12))


def kdp_inside_margin(page_count: int) -> float:
    pages = max(1, int(page_count or 1))
    if pages <= 150:
        return 0.375
    if pages <= 300:
        return 0.5
    if pages <= 500:
        return 0.625
    if pages <= 700:
        return 0.75
    return 0.875


def kdp_profile(
    *,
    trim_code: str = "6x9",
    page_count: int = 150,
    bleed: bool = False,
) -> dict[str, Any]:
    if trim_code not in KDP_TRIM_SIZES:
        raise ValueError("Unsupported KDP trim size.")
    width, height = KDP_TRIM_SIZES[trim_code]
    inside = kdp_inside_margin(page_count)
    outside = 0.375
    top_bottom = 1.0 if bleed else 0.75
    maximum = KDP_BW_WHITE_MAX_PAGES[trim_code]
    warnings: list[str] = []
    if int(page_count) > maximum:
        warnings.append(
            f"Estimated page count exceeds the {maximum}-page black-ink/white-paper "
            f"limit published for {trim_code}. Other paper/ink combinations can be lower."
        )
    return {
        "trimCode": trim_code,
        "trimWidth": width,
        "trimHeight": height,
        "pageWidth": width + (0.125 if bleed else 0.0),
        "pageHeight": height + (0.25 if bleed else 0.0),
        "insideMargin": inside,
        "outsideMargin": outside,
        "topMargin": top_bottom,
        "bottomMargin": top_bottom,
        "bleed": bool(bleed),
        "pageCount": int(page_count),
        "blackInkWhitePaperMaxPages": maximum,
        "warnings": warnings,
    }


@dataclass(slots=True)
class ManuscriptError(RuntimeError):
    message: str
    code: str = "MANUSCRIPT_ERROR"
    status_code: int = 400

    def __post_init__(self) -> None:
        RuntimeError.__init__(self, self.message)


class ManuscriptService:
    def __init__(
        self,
        db: SupabaseDB,
        ai: AIService,
        projects: ProjectService,
    ) -> None:
        self.db = db
        self.ai = ai
        self.projects = projects

    async def list(self, *, user_id: str, project_id: str) -> list[dict[str, Any]]:
        project = await self.projects.get(user_id, project_id)
        return await self.db.select(
            "manuscripts",
            filters={
                "user_id": eq(user_id),
                "project_id": eq(project["id"]),
                "archived_at": "is.null",
            },
            order="updated_at.desc",
            limit=100,
        )

    async def get(self, *, user_id: str, manuscript_id: str) -> dict[str, Any]:
        try:
            normalized = normalize_chat_id(manuscript_id)
        except Exception as exc:
            raise ManuscriptError("Manuscript not found.", "MANUSCRIPT_NOT_FOUND", 404) from exc
        row = await self.db.select_one(
            "manuscripts",
            filters={
                "id": eq(normalized),
                "user_id": eq(user_id),
                "archived_at": "is.null",
            },
        )
        if not row:
            raise ManuscriptError("Manuscript not found.", "MANUSCRIPT_NOT_FOUND", 404)
        return row

    async def create(
        self,
        *,
        user_id: str,
        project_id: str,
        title: str,
        subtitle: str = "",
        author_name: str = "",
        trim_code: str = "6x9",
        bleed: bool = False,
        metadata: dict[str, Any] | None = None,
    ) -> dict[str, Any]:
        project = await self.projects.get(user_id, project_id)
        title = _clean(title, 180)
        if not title:
            raise ManuscriptError("A manuscript title is required.")
        if trim_code not in KDP_TRIM_SIZES:
            raise ManuscriptError("Choose a supported KDP trim size.", "INVALID_TRIM_SIZE")
        width, height = KDP_TRIM_SIZES[trim_code]
        safe_metadata = dict(metadata or {})
        if "targetWords" in safe_metadata:
            try:
                safe_metadata["targetWords"] = max(
                    MIN_TARGET_WORDS,
                    min(MAX_TARGET_WORDS, int(safe_metadata["targetWords"])),
                )
            except (TypeError, ValueError):
                safe_metadata["targetWords"] = DEFAULT_TARGET_WORDS
        if "premise" in safe_metadata:
            safe_metadata["premise"] = _clean(safe_metadata["premise"], 1200)
        row = {
            "id": str(uuid4()),
            "user_id": user_id,
            "project_id": project["id"],
            "title": title,
            "subtitle": _clean(subtitle, 240),
            "author_name": _clean(author_name, 160),
            "trim_code": trim_code,
            "trim_width": width,
            "trim_height": height,
            "bleed": bool(bleed),
            "status": "draft",
            "metadata": safe_metadata,
            "updated_at": _now(),
        }
        result = await self.db.insert("manuscripts", row)
        return (result or [row])[0]

    async def list_sections(self, *, user_id: str, manuscript_id: str) -> list[dict[str, Any]]:
        manuscript = await self.get(user_id=user_id, manuscript_id=manuscript_id)
        return await self.db.select(
            "manuscript_sections",
            filters={"user_id": eq(user_id), "manuscript_id": eq(manuscript["id"])},
            order="position.asc",
            limit=500,
        )

    async def add_section(
        self,
        *,
        user_id: str,
        manuscript_id: str,
        title: str,
        section_type: str = "chapter",
        content: str = "",
    ) -> dict[str, Any]:
        manuscript = await self.get(user_id=user_id, manuscript_id=manuscript_id)
        title = _clean(title, 180)
        if not title:
            raise ManuscriptError("A section title is required.")
        section_type = str(section_type or "chapter").lower()
        if section_type not in {"frontmatter", "chapter", "scene", "backmatter"}:
            raise ManuscriptError("Invalid manuscript section type.", "INVALID_SECTION_TYPE")
        sections = await self.list_sections(user_id=user_id, manuscript_id=manuscript["id"])
        text = str(content or "").strip()[:250000]
        row = {
            "id": str(uuid4()),
            "user_id": user_id,
            "manuscript_id": manuscript["id"],
            "section_type": section_type,
            "title": title,
            "position": len(sections) + 1,
            "content": text,
            "summary": self._summary(text),
            "word_count": word_count(text),
            "status": "draft" if text else "outline",
            "updated_at": _now(),
        }
        result = await self.db.insert("manuscript_sections", row)
        await self._touch(manuscript["id"], user_id)
        return (result or [row])[0]

    async def update_section(
        self,
        *,
        user_id: str,
        manuscript_id: str,
        section_id: str,
        changes: dict[str, Any],
    ) -> dict[str, Any]:
        manuscript = await self.get(user_id=user_id, manuscript_id=manuscript_id)
        normalized = normalize_chat_id(section_id)
        row = await self.db.select_one(
            "manuscript_sections",
            filters={
                "id": eq(normalized),
                "user_id": eq(user_id),
                "manuscript_id": eq(manuscript["id"]),
            },
        )
        if not row:
            raise ManuscriptError("Section not found.", "SECTION_NOT_FOUND", 404)
        updates: dict[str, Any] = {"updated_at": _now()}
        if "title" in changes:
            title = _clean(changes.get("title"), 180)
            if title:
                updates["title"] = title
        if "content" in changes:
            text = str(changes.get("content") or "").strip()[:250000]
            updates.update(
                {
                    "content": text,
                    "summary": self._summary(text),
                    "word_count": word_count(text),
                    "status": "draft" if text else "outline",
                }
            )
        if "status" in changes and str(changes["status"]) in {"outline", "draft", "revised", "final"}:
            updates["status"] = str(changes["status"])
        result = await self.db.update(
            "manuscript_sections",
            updates,
            filters={"id": eq(normalized), "user_id": eq(user_id)},
        )
        await self._touch(manuscript["id"], user_id)
        return (result or [{**row, **updates}])[0]

    @staticmethod
    def _decode_blueprint(text: str) -> dict[str, Any]:
        clean = str(text or "").strip()
        clean = re.sub(r"^```(?:json)?\s*", "", clean, flags=re.I)
        clean = re.sub(r"\s*```$", "", clean)
        start = clean.find("{")
        end = clean.rfind("}")
        if start < 0 or end <= start:
            return {}
        try:
            value = json.loads(clean[start:end + 1])
        except (json.JSONDecodeError, TypeError, ValueError):
            return {}
        return value if isinstance(value, dict) else {}

    @classmethod
    def _normalize_blueprint(
        cls,
        value: dict[str, Any],
        *,
        brief: str,
        preferred_title: str,
        target_words: int,
        chapter_count: int,
    ) -> dict[str, Any]:
        title = _clean(preferred_title, 180) or _clean(value.get("title"), 180)
        title = title or title_from_prompt(brief)
        subtitle = _clean(value.get("subtitle"), 240)
        genre = _clean(value.get("genre"), 120)
        premise = _clean(value.get("premise"), 1200) or _clean(brief, 1200)
        try:
            proposed_target = int(str(value.get("targetWords") or target_words).replace(",", ""))
        except (TypeError, ValueError):
            proposed_target = target_words
        normalized_target = max(MIN_TARGET_WORDS, min(MAX_TARGET_WORDS, proposed_target))

        raw_sections = value.get("chapters") or value.get("sections") or []
        sections: list[dict[str, str]] = []
        if isinstance(raw_sections, list):
            for index, item in enumerate(raw_sections[:80], start=1):
                if not isinstance(item, dict):
                    continue
                chapter_title = _clean(item.get("title"), 180) or f"Chapter {index}"
                purpose = _clean(
                    item.get("purpose") or item.get("brief") or item.get("summary"),
                    1800,
                )
                sections.append(
                    {
                        "title": chapter_title,
                        "purpose": purpose
                        or f"Advance the central conflict and continuity established for {title}.",
                    }
                )

        if len(sections) < 8:
            sections = [
                {
                    "title": f"Chapter {index}",
                    "purpose": (
                        f"Develop the next necessary movement of {title}, preserving the user's brief, "
                        "causal continuity, and forward momentum."
                    ),
                }
                for index in range(1, chapter_count + 1)
            ]

        return {
            "title": title,
            "subtitle": subtitle,
            "genre": genre,
            "premise": premise,
            "targetWords": normalized_target,
            "chapters": sections,
        }

    async def plan_blueprint(
        self,
        *,
        user: dict[str, Any],
        brief: str,
        target_words: int | None = None,
        chapter_count: int | None = None,
        preferred_title: str = "",
        project_id: str | None = None,
    ) -> tuple[dict[str, Any], dict[str, Any]]:
        clean_brief = str(brief or "").strip()[:12000]
        if not clean_brief:
            raise ManuscriptError("Describe the manuscript Crump should plan.", "MANUSCRIPT_BRIEF_REQUIRED")
        target = target_words_from_prompt(clean_brief, target_words or DEFAULT_TARGET_WORDS)
        chapters = chapter_count_from_prompt(clean_brief, chapter_count or DEFAULT_CHAPTER_COUNT)
        project_context: dict[str, Any] = {}
        if project_id:
            project_context = await self.projects.hydrate_context(user["id"], project_id)

        prompt = f"""Design a complete, original, book-length manuscript blueprint from the user's brief.
Return JSON only—no Markdown fence and no commentary—with exactly this shape:
{{
  "title": "...",
  "subtitle": "...",
  "genre": "...",
  "premise": "...",
  "targetWords": {target},
  "chapters": [{{"title": "...", "purpose": "1-3 specific sentences"}}]
}}

Create approximately {chapters} chapters. Every chapter must cause the next one: identify its
dramatic or argumentative purpose, important development, and continuity burden. Avoid repetitive
beats, filler chapters, generic titles, and premature resolution. Preserve supplied Project canon.
The plan must be substantial enough to support roughly {target:,} finished words.

User brief:
{clean_brief}"""
        result = await self.ai.chat(
            {
                "message": prompt,
                "history": [],
                "assistantName": "Crump",
                "user": {"name": user.get("full_name") or user.get("name") or ""},
                "relevantContext": project_context,
                "workMode": "work",
            }
        )
        decoded = self._decode_blueprint(str(result.get("response") or ""))
        blueprint = self._normalize_blueprint(
            decoded,
            brief=clean_brief,
            preferred_title=preferred_title,
            target_words=target,
            chapter_count=chapters,
        )
        return blueprint, result

    async def _apply_blueprint_rows(
        self,
        *,
        user_id: str,
        manuscript: dict[str, Any],
        blueprint: dict[str, Any],
        replace_outlines: bool = False,
    ) -> list[dict[str, Any]]:
        existing = await self.list_sections(user_id=user_id, manuscript_id=manuscript["id"])
        if any(str(item.get("content") or "").strip() for item in existing):
            raise ManuscriptError(
                "This manuscript already has drafted text. Crump will not erase it to replace the plan.",
                "MANUSCRIPT_HAS_DRAFTS",
                409,
            )
        if existing and not replace_outlines:
            raise ManuscriptError(
                "This manuscript already has a chapter plan.",
                "MANUSCRIPT_ALREADY_PLANNED",
                409,
            )
        if existing:
            await self.db.delete(
                "manuscript_sections",
                filters={"user_id": eq(user_id), "manuscript_id": eq(manuscript["id"])},
            )

        now = _now()
        rows = [
            {
                "id": str(uuid4()),
                "user_id": user_id,
                "manuscript_id": manuscript["id"],
                "section_type": "chapter",
                "title": item["title"],
                "position": index,
                "content": "",
                "summary": item["purpose"],
                "word_count": 0,
                "status": "outline",
                "updated_at": now,
            }
            for index, item in enumerate(blueprint["chapters"], start=1)
        ]
        inserted = await self.db.insert("manuscript_sections", rows)
        metadata = dict(manuscript.get("metadata") or {})
        metadata.update(
            {
                "genre": blueprint.get("genre") or "",
                "premise": blueprint.get("premise") or "",
                "targetWords": int(blueprint.get("targetWords") or DEFAULT_TARGET_WORDS),
                "plannedChapterCount": len(rows),
                "blueprintCreatedAt": now,
            }
        )
        updates: dict[str, Any] = {"metadata": metadata, "updated_at": now}
        if blueprint.get("subtitle") and not manuscript.get("subtitle"):
            updates["subtitle"] = blueprint["subtitle"]
        await self.db.update(
            "manuscripts",
            updates,
            filters={"id": eq(manuscript["id"]), "user_id": eq(user_id)},
        )
        return inserted or rows

    async def apply_blueprint(
        self,
        *,
        user: dict[str, Any],
        manuscript_id: str,
        brief: str,
        target_words: int | None = None,
        chapter_count: int | None = None,
        replace_outlines: bool = False,
    ) -> dict[str, Any]:
        manuscript = await self.get(user_id=user["id"], manuscript_id=manuscript_id)
        blueprint, result = await self.plan_blueprint(
            user=user,
            brief=brief,
            target_words=target_words,
            chapter_count=chapter_count,
            preferred_title=str(manuscript.get("title") or ""),
            project_id=str(manuscript["project_id"]),
        )
        sections = await self._apply_blueprint_rows(
            user_id=user["id"],
            manuscript=manuscript,
            blueprint=blueprint,
            replace_outlines=replace_outlines,
        )
        return {"blueprint": blueprint, "sections": sections, "model": result.get("model")}

    async def begin_long_form(
        self,
        *,
        user: dict[str, Any],
        brief: str,
        project_id: str | None = None,
        chat_id: str | None = None,
        preferred_format: str = "docx",
        project_limit: int = 2,
    ) -> dict[str, Any]:
        # Reject a full workspace before paying for its blueprint when the user
        # has no Project slot available. Existing Projects are validated by the
        # context hydration inside plan_blueprint.
        if not project_id and await self.projects.count(user["id"]) >= project_limit:
            raise ManuscriptError(
                "Choose an existing Project before starting another long-form workspace.",
                "PROJECT_LIMIT_REACHED",
                403,
            )
        blueprint, model_result = await self.plan_blueprint(
            user=user,
            brief=brief,
            project_id=project_id,
        )
        created_project = False
        if project_id:
            project = await self.projects.get(user["id"], project_id)
        else:
            project = await self.projects.create(
                user_id=user["id"],
                name=blueprint["title"],
                description=blueprint.get("premise") or "Long-form manuscript workspace",
                instructions=f"Original long-form brief:\n{str(brief or '').strip()[:10000]}",
            )
            created_project = True

        manuscript = await self.create(
            user_id=user["id"],
            project_id=project["id"],
            title=blueprint["title"],
            subtitle=blueprint.get("subtitle") or "",
            author_name=user.get("full_name") or user.get("name") or "",
            metadata={
                "preferredExportFormat": preferred_format,
                "source": "chat_long_form_handoff",
            },
        )
        sections = await self._apply_blueprint_rows(
            user_id=user["id"],
            manuscript=manuscript,
            blueprint=blueprint,
        )
        if chat_id:
            await self.projects.attach_chat(
                user_id=user["id"],
                project_id=project["id"],
                chat_id=chat_id,
            )
        return {
            "response": (
                f"I created a persistent Manuscript workspace for **{blueprint['title']}** with "
                f"{len(sections)} planned chapters and a {int(blueprint['targetWords']):,}-word target. "
                "I moved it out of the chat box so the outline, chapter drafts, Project canon, revisions, "
                "and exports can survive across sessions. Open the workspace below to review the plan or "
                "draft the next chapter with Crump."
            ),
            "model": model_result.get("model"),
            "usage": model_result.get("usage") or {},
            "stopReason": model_result.get("stopReason"),
            "projectId": project["id"],
            "manuscriptWorkspace": {
                "projectId": project["id"],
                "projectName": project.get("name"),
                "projectCreated": created_project,
                "manuscriptId": manuscript["id"],
                "title": blueprint["title"],
                "chapterCount": len(sections),
                "targetWords": int(blueprint["targetWords"]),
                "preferredExportFormat": preferred_format,
            },
        }

    @staticmethod
    def progress(manuscript: dict[str, Any], sections: list[dict[str, Any]]) -> dict[str, Any]:
        metadata = manuscript.get("metadata") if isinstance(manuscript.get("metadata"), dict) else {}
        total_words = sum(int(item.get("word_count") or word_count(item.get("content") or "")) for item in sections)
        try:
            target_words = max(1, int(metadata.get("targetWords") or DEFAULT_TARGET_WORDS))
        except (TypeError, ValueError):
            target_words = DEFAULT_TARGET_WORDS
        drafted = sum(1 for item in sections if str(item.get("content") or "").strip())
        next_section = next((item for item in sections if not str(item.get("content") or "").strip()), None)
        return {
            "wordCount": total_words,
            "targetWords": target_words,
            "wordProgress": min(100, round(total_words * 100 / target_words, 1)),
            "draftedSections": drafted,
            "plannedSections": len(sections),
            "nextSectionId": next_section.get("id") if next_section else None,
            "nextSectionTitle": next_section.get("title") if next_section else None,
            "complete": bool(sections) and drafted == len(sections),
        }

    async def draft_next(
        self,
        *,
        user: dict[str, Any],
        manuscript_id: str,
        instruction: str = "",
    ) -> dict[str, Any]:
        sections = await self.list_sections(user_id=user["id"], manuscript_id=manuscript_id)
        target = next((item for item in sections if not str(item.get("content") or "").strip()), None)
        if not target:
            raise ManuscriptError(
                "Every planned section already has a draft.",
                "MANUSCRIPT_DRAFT_COMPLETE",
                409,
            )
        return await self.draft_section(
            user=user,
            manuscript_id=manuscript_id,
            section_id=str(target["id"]),
            instruction=instruction,
        )

    async def draft_section(
        self,
        *,
        user: dict[str, Any],
        manuscript_id: str,
        section_id: str,
        instruction: str = "",
    ) -> dict[str, Any]:
        manuscript = await self.get(user_id=user["id"], manuscript_id=manuscript_id)
        sections = await self.list_sections(user_id=user["id"], manuscript_id=manuscript["id"])
        target_id = normalize_chat_id(section_id)
        index = next((i for i, item in enumerate(sections) if item.get("id") == target_id), -1)
        if index < 0:
            raise ManuscriptError("Section not found.", "SECTION_NOT_FOUND", 404)
        target = sections[index]
        project_context = await self.projects.hydrate_context(user["id"], manuscript["project_id"])
        previous = sections[max(0, index - 3):index]
        upcoming = sections[index + 1:index + 4]
        immediate_previous = previous[-1] if previous else {}
        canon_rows = (project_context.get("canon") or [])[:12]
        compact_canon = [
            {
                "kind": item.get("kind"),
                "label": item.get("label"),
                "content": str(item.get("content") or "")[:900],
            }
            for item in canon_rows
        ]
        project = project_context.get("project") or {}
        metadata = manuscript.get("metadata") if isinstance(manuscript.get("metadata"), dict) else {}
        try:
            manuscript_target = int(metadata.get("targetWords") or DEFAULT_TARGET_WORDS)
        except (TypeError, ValueError):
            manuscript_target = DEFAULT_TARGET_WORDS
        chapter_target = max(
            1_800,
            min(
                4_500,
                round(manuscript_target / max(1, len(sections))),
            ),
        )
        context = {
            "sectionBrief": str(target.get("summary") or "")[:1800],
            "currentDraft": str(target.get("content") or "")[-4500:],
            "previousEnding": str(immediate_previous.get("content") or "")[-3200:],
            "previousSections": [
                {"title": item.get("title"), "continuity": str(item.get("summary") or "")[:900]}
                for item in previous
            ],
            "upcomingSections": [
                {"title": item.get("title"), "brief": str(item.get("summary") or "")[:500]}
                for item in upcoming
            ],
            "manuscript": {
                "title": manuscript.get("title"),
                "subtitle": manuscript.get("subtitle"),
                "author": manuscript.get("author_name"),
                "genre": metadata.get("genre"),
                "premise": metadata.get("premise"),
                "targetWords": metadata.get("targetWords"),
            },
            "project": {
                "name": project.get("name"),
                "description": str(project.get("description") or "")[:900],
                "instructions": str(project.get("instructions") or "")[:2200],
                "canon": compact_canon,
            },
        }
        direction = str(instruction or "").strip()[:5000]
        action = "Revise and expand" if str(target.get("content") or "").strip() else "Draft"
        prompt = (
            f"{action} the manuscript section titled {target.get('title')!r}. "
            "Write polished long-form prose, not an outline or summary. Fulfill the supplied section brief. "
            "Continue causally from the previous ending and maintain Project canon, chronology, voice, POV, "
            "character knowledge, and unresolved threads. Prefer concrete scenes, sensory detail, deliberate "
            "pacing, and earned transitions. Do not add commentary before or after the manuscript text. "
            f"Target roughly {chapter_target:,} words unless the author clearly requests another length."
        )
        if direction:
            prompt += f"\n\nSpecific direction from the author:\n{direction}"
        result = await self.ai.chat(
            {
                "message": prompt,
                "history": [],
                "assistantName": "Crump",
                "user": {"name": user.get("full_name") or user.get("name") or ""},
                "relevantContext": context,
                "workMode": "work",
            }
        )
        text = str(result.get("response") or "").strip()
        if not text:
            raise ManuscriptError("Crump returned an empty manuscript draft.", "EMPTY_DRAFT", 502)
        return await self.update_section(
            user_id=user["id"],
            manuscript_id=manuscript["id"],
            section_id=target_id,
            changes={"content": text},
        )

    async def export(
        self,
        *,
        user_id: str,
        manuscript_id: str,
        export_format: str,
    ) -> tuple[bytes, str, str, dict[str, Any]]:
        manuscript = await self.get(user_id=user_id, manuscript_id=manuscript_id)
        sections = await self.list_sections(user_id=user_id, manuscript_id=manuscript["id"])
        if not sections:
            raise ManuscriptError("Add at least one manuscript section before exporting.")
        total_words = sum(int(item.get("word_count") or word_count(item.get("content") or "")) for item in sections)
        page_count = estimate_pages(total_words)
        profile = kdp_profile(
            trim_code=str(manuscript.get("trim_code") or "6x9"),
            page_count=page_count,
            bleed=bool(manuscript.get("bleed")),
        )
        fmt = str(export_format or "docx").lower().strip(".")
        safe_title = re.sub(r"[^A-Za-z0-9._-]+", "_", str(manuscript.get("title") or "manuscript")).strip("_")
        safe_title = safe_title[:90] or "manuscript"
        metadata = {
            "kdpProfile": profile,
            "estimatedPageCount": page_count,
            "wordCount": total_words,
            "manuscriptId": manuscript["id"],
            "projectId": manuscript["project_id"],
        }
        if fmt == "docx":
            return (
                self._docx(manuscript, sections, profile),
                f"{safe_title}_KDP.docx",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                metadata,
            )
        if fmt == "pdf":
            return (
                self._pdf(manuscript, sections, profile),
                f"{safe_title}_KDP.pdf",
                "application/pdf",
                metadata,
            )
        if fmt == "epub":
            return (
                self._epub(manuscript, sections),
                f"{safe_title}.epub",
                "application/epub+zip",
                metadata,
            )
        raise ManuscriptError("Export format must be DOCX, PDF, or EPUB.", "INVALID_EXPORT_FORMAT")

    async def _touch(self, manuscript_id: str, user_id: str) -> None:
        await self.db.update(
            "manuscripts",
            {"updated_at": _now()},
            filters={"id": eq(manuscript_id), "user_id": eq(user_id)},
        )

    @staticmethod
    def _summary(text: str) -> str:
        compact = " ".join(str(text or "").split())
        if len(compact) <= 1800:
            return compact
        # Preserve both setup and ending. The final state of a chapter is usually
        # more important to continuity than another slice from its opening pages.
        return f"{compact[:850]} … [chapter ending] … {compact[-850:]}"

    @staticmethod
    def _add_page_number(paragraph: Any) -> None:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        begin = OxmlElement("w:fldChar")
        begin.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = " PAGE "
        end = OxmlElement("w:fldChar")
        end.set(qn("w:fldCharType"), "end")
        run._r.extend([begin, instr, end])

    @staticmethod
    def _set_mirror_margins(document: Document) -> None:
        settings = document.settings._element
        if settings.find(qn("w:mirrorMargins")) is None:
            settings.append(OxmlElement("w:mirrorMargins"))

    def _docx(
        self,
        manuscript: dict[str, Any],
        sections: list[dict[str, Any]],
        profile: dict[str, Any],
    ) -> bytes:
        document = Document()
        section = document.sections[0]
        section.start_type = WD_SECTION.NEW_PAGE
        section.page_width = Inches(profile["pageWidth"])
        section.page_height = Inches(profile["pageHeight"])
        section.top_margin = Inches(profile["topMargin"])
        section.bottom_margin = Inches(profile["bottomMargin"])
        section.left_margin = Inches(profile["insideMargin"])
        section.right_margin = Inches(profile["outsideMargin"])
        self._set_mirror_margins(document)

        normal = document.styles["Normal"]
        normal.font.name = "Garamond"
        normal.font.size = Pt(11)
        normal.paragraph_format.line_spacing = 1.15
        normal.paragraph_format.space_after = Pt(0)
        normal.paragraph_format.first_line_indent = Inches(0.22)

        title = document.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.paragraph_format.space_before = Inches(2.1)
        title_run = title.add_run(str(manuscript.get("title") or "Untitled"))
        title_run.bold = True
        title_run.font.name = "Garamond"
        title_run.font.size = Pt(26)
        if manuscript.get("subtitle"):
            subtitle = document.add_paragraph()
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
            subtitle.add_run(str(manuscript["subtitle"])).font.size = Pt(14)
        if manuscript.get("author_name"):
            author = document.add_paragraph()
            author.alignment = WD_ALIGN_PARAGRAPH.CENTER
            author.paragraph_format.space_before = Inches(1.2)
            author.add_run(str(manuscript["author_name"])).font.size = Pt(12)
        document.add_page_break()

        copyright_text = (
            f"Copyright © {datetime.now().year} {manuscript.get('author_name') or 'Author'}. "
            "All rights reserved."
        )
        copyright_paragraph = document.add_paragraph(copyright_text)
        copyright_paragraph.paragraph_format.space_before = Inches(2.0)
        copyright_paragraph.paragraph_format.first_line_indent = Inches(0)
        copyright_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Front matter stays unnumbered. Begin the book body in a new section and
        # restart Arabic page numbering at 1 for a cleaner trade-book interior.
        body_section = document.add_section(WD_SECTION.NEW_PAGE)
        body_section.page_width = Inches(profile["pageWidth"])
        body_section.page_height = Inches(profile["pageHeight"])
        body_section.top_margin = Inches(profile["topMargin"])
        body_section.bottom_margin = Inches(profile["bottomMargin"])
        body_section.left_margin = Inches(profile["insideMargin"])
        body_section.right_margin = Inches(profile["outsideMargin"])
        body_section.footer.is_linked_to_previous = False
        self._add_page_number(body_section.footer.paragraphs[0])
        sect_pr = body_section._sectPr
        pg_num = sect_pr.find(qn("w:pgNumType"))
        if pg_num is None:
            pg_num = OxmlElement("w:pgNumType")
            sect_pr.append(pg_num)
        pg_num.set(qn("w:start"), "1")

        for index, item in enumerate(sections):
            if index:
                document.add_page_break()
            heading = document.add_paragraph()
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            heading.paragraph_format.space_before = Inches(0.8)
            heading.paragraph_format.space_after = Inches(0.5)
            run = heading.add_run(str(item.get("title") or "Chapter"))
            run.bold = True
            run.font.name = "Garamond"
            run.font.size = Pt(18)
            self._append_docx_content(document, str(item.get("content") or ""))

        stream = BytesIO()
        document.save(stream)
        return stream.getvalue()

    @staticmethod
    def _append_docx_content(document: Document, content: str) -> None:
        blocks = re.split(r"\n\s*\n", content.strip()) if content.strip() else []
        for block in blocks:
            text = block.strip()
            if not text:
                continue
            if text in {"***", "* * *", "⸻", "---"}:
                divider = document.add_paragraph("⸻")
                divider.alignment = WD_ALIGN_PARAGRAPH.CENTER
                divider.paragraph_format.first_line_indent = Inches(0)
                divider.paragraph_format.space_before = Pt(10)
                divider.paragraph_format.space_after = Pt(10)
                continue
            text = re.sub(r"^#{1,6}\s+", "", text)
            paragraph = document.add_paragraph(text)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    def _pdf(
        self,
        manuscript: dict[str, Any],
        sections: list[dict[str, Any]],
        profile: dict[str, Any],
    ) -> bytes:
        stream = BytesIO()
        page_size = (profile["pageWidth"] * inch, profile["pageHeight"] * inch)
        doc = SimpleDocTemplate(
            stream,
            pagesize=page_size,
            leftMargin=max(profile["insideMargin"], profile["outsideMargin"]) * inch,
            # Use the larger gutter on both sides in PDF. This is intentionally
            # conservative: a fixed PDF cannot rely on a word processor's mirror-margin
            # setting, so both inside and outside edges remain at or above KDP's
            # required minimum regardless of left/right page parity.
            rightMargin=max(profile["insideMargin"], profile["outsideMargin"]) * inch,
            topMargin=profile["topMargin"] * inch,
            bottomMargin=profile["bottomMargin"] * inch,
            title=str(manuscript.get("title") or "Untitled"),
            author=str(manuscript.get("author_name") or ""),
        )
        styles = getSampleStyleSheet()
        body = ParagraphStyle(
            "CrumpBookBody",
            parent=styles["BodyText"],
            fontName="Times-Roman",
            fontSize=10.8,
            leading=15.0,
            alignment=TA_JUSTIFY,
            firstLineIndent=16,
            spaceAfter=7,
        )
        chapter = ParagraphStyle(
            "CrumpChapter",
            parent=styles["Heading1"],
            fontName="Times-Bold",
            fontSize=18,
            leading=22,
            alignment=TA_CENTER,
            spaceBefore=46,
            spaceAfter=34,
        )
        title_style = ParagraphStyle(
            "CrumpTitle",
            parent=styles["Title"],
            fontName="Times-Bold",
            fontSize=26,
            leading=30,
            alignment=TA_CENTER,
            spaceBefore=120,
            spaceAfter=20,
        )
        centered = ParagraphStyle(
            "CrumpCentered",
            parent=body,
            alignment=TA_CENTER,
            firstLineIndent=0,
            spaceBefore=10,
            spaceAfter=10,
        )
        story: list[Any] = [Paragraph(html.escape(str(manuscript.get("title") or "Untitled")), title_style)]
        if manuscript.get("subtitle"):
            story.append(Paragraph(html.escape(str(manuscript["subtitle"])), centered))
        story.append(Spacer(1, 52))
        if manuscript.get("author_name"):
            story.append(Paragraph(html.escape(str(manuscript["author_name"])), centered))
        story.append(PageBreak())
        story.append(Spacer(1, 120))
        story.append(
            Paragraph(
                html.escape(
                    f"Copyright © {datetime.now().year} "
                    f"{manuscript.get('author_name') or 'Author'}. All rights reserved."
                ),
                body,
            )
        )
        story.append(PageBreak())
        for index, item in enumerate(sections):
            if index:
                story.append(PageBreak())
            story.append(Paragraph(html.escape(str(item.get("title") or "Chapter")), chapter))
            for block in re.split(r"\n\s*\n", str(item.get("content") or "").strip()):
                text = block.strip()
                if not text:
                    continue
                if text in {"***", "* * *", "⸻", "---"}:
                    story.append(Paragraph("— • —", centered))
                    continue
                text = re.sub(r"^#{1,6}\s+", "", text)
                escaped = html.escape(text).replace("\n", "<br/>")
                story.append(Paragraph(escaped, body))

        def footer(canvas: Any, document: Any) -> None:
            # Title and copyright pages are unnumbered; body begins at page 1.
            if document.page <= 2:
                return
            canvas.saveState()
            canvas.setFont("Times-Roman", 8)
            canvas.setFillColorRGB(0.35, 0.35, 0.35)
            canvas.drawCentredString(page_size[0] / 2, 0.42 * inch, str(document.page - 2))
            canvas.restoreState()

        doc.build(story, onFirstPage=footer, onLaterPages=footer)
        return stream.getvalue()

    def _epub(self, manuscript: dict[str, Any], sections: list[dict[str, Any]]) -> bytes:
        stream = BytesIO()
        title = str(manuscript.get("title") or "Untitled")
        author = str(manuscript.get("author_name") or "")
        book_id = f"urn:uuid:{manuscript['id']}"
        items: list[tuple[str, str, str]] = []
        nav_links: list[str] = []
        spine: list[str] = []
        for index, section in enumerate(sections, start=1):
            item_id = f"chapter{index}"
            filename = f"{item_id}.xhtml"
            section_title = str(section.get("title") or f"Chapter {index}")
            paragraphs = []
            for block in re.split(r"\n\s*\n", str(section.get("content") or "").strip()):
                text = block.strip()
                if not text:
                    continue
                if text in {"***", "* * *", "⸻", "---"}:
                    paragraphs.append('<p class="scene">⸻</p>')
                else:
                    text = re.sub(r"^#{1,6}\s+", "", text)
                    paragraphs.append(f"<p>{html.escape(text).replace(chr(10), '<br/>')}</p>")
            xhtml = (
                '<?xml version="1.0" encoding="utf-8"?>'
                '<html xmlns="http://www.w3.org/1999/xhtml"><head>'
                f"<title>{html.escape(section_title)}</title>"
                '<link rel="stylesheet" type="text/css" href="style.css"/></head><body>'
                f"<h1>{html.escape(section_title)}</h1>{''.join(paragraphs)}</body></html>"
            )
            items.append((item_id, filename, xhtml))
            nav_links.append(f'<li><a href="{filename}">{html.escape(section_title)}</a></li>')
            spine.append(f'<itemref idref="{item_id}"/>')

        nav = (
            '<?xml version="1.0" encoding="utf-8"?>'
            '<html xmlns="http://www.w3.org/1999/xhtml" '
            'xmlns:epub="http://www.idpf.org/2007/ops"><head><title>Contents</title></head>'
            f"<body><nav epub:type=\"toc\"><h1>Contents</h1><ol>{''.join(nav_links)}</ol></nav></body></html>"
        )
        manifest = "".join(
            f'<item id="{item_id}" href="{filename}" media-type="application/xhtml+xml"/>'
            for item_id, filename, _ in items
        )
        package = (
            '<?xml version="1.0" encoding="UTF-8"?>'
            '<package xmlns="http://www.idpf.org/2007/opf" version="3.0" unique-identifier="bookid">'
            '<metadata xmlns:dc="http://purl.org/dc/elements/1.1/">'
            f'<dc:identifier id="bookid">{html.escape(book_id)}</dc:identifier>'
            f"<dc:title>{html.escape(title)}</dc:title>"
            f"<dc:creator>{html.escape(author)}</dc:creator>"
            '<dc:language>en</dc:language></metadata><manifest>'
            '<item id="nav" href="nav.xhtml" media-type="application/xhtml+xml" properties="nav"/>'
            '<item id="css" href="style.css" media-type="text/css"/>'
            f"{manifest}</manifest><spine>{''.join(spine)}</spine></package>"
        )
        container_xml = (
            '<?xml version="1.0" encoding="UTF-8"?>'
            '<container version="1.0" xmlns="urn:oasis:names:tc:opendocument:xmlns:container">'
            '<rootfiles><rootfile full-path="OEBPS/package.opf" '
            'media-type="application/oebps-package+xml"/></rootfiles></container>'
        )
        css = (
            "body{font-family:serif;line-height:1.5;margin:5%;}"
            "h1{text-align:center;page-break-before:always;margin:18% 0 8%;}"
            "p{text-indent:1.2em;margin:.35em 0;text-align:justify;}"
            ".scene{text-indent:0;text-align:center;margin:1.4em 0;}"
        )
        with zipfile.ZipFile(stream, "w") as archive:
            archive.writestr("mimetype", "application/epub+zip", compress_type=zipfile.ZIP_STORED)
            archive.writestr("META-INF/container.xml", container_xml)
            archive.writestr("OEBPS/nav.xhtml", nav)
            archive.writestr("OEBPS/package.opf", package)
            archive.writestr("OEBPS/style.css", css)
            for _, filename, xhtml in items:
                archive.writestr(f"OEBPS/{filename}", xhtml)
        return stream.getvalue()
